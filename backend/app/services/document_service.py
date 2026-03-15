"""
Document Generation Service
Generates professional Word (.docx) and PDF documents for each tender.
"""

import logging
import os
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
)

logger = logging.getLogger(__name__)

DOCS_DIR = Path("generated_docs")
DOCS_DIR.mkdir(exist_ok=True)

# ── Brand colours ─────────────────────────────────────────────────────────────
BRAND_DARK_BLUE = RGBColor(0x0D, 0x47, 0xA1)   # #0D47A1
BRAND_GREEN = RGBColor(0x1B, 0x5E, 0x20)        # #1B5E20
BRAND_GOLD = RGBColor(0xF5, 0x7F, 0x17)         # #F57F17
RL_DARK_BLUE = colors.HexColor("#0D47A1")
RL_GREEN = colors.HexColor("#1B5E20")
RL_GOLD = colors.HexColor("#F57F17")
RL_LIGHT_GREY = colors.HexColor("#F5F5F5")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _format_value(v: Optional[float]) -> str:
    if v is None:
        return "Not specified"
    return f"PKR {v:,.0f}"


def _format_date(d: Optional[datetime]) -> str:
    if d is None:
        return "Not specified"
    return d.strftime("%d %B %Y")


# ── Word Document ─────────────────────────────────────────────────────────────

def generate_word_document(tender: Dict[str, Any]) -> Tuple[bytes, str]:
    """
    Generate a branded Word document for a tender.
    Returns (bytes, filename).
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover header ──────────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("🏛 TenderIQ Pakistan")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK_BLUE

    sub_run = doc.add_paragraph()
    sub_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_run.add_run("Government Procurement Intelligence Report")
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND_GOLD
    r.font.italic = True

    doc.add_paragraph()

    # Horizontal rule
    hr = doc.add_paragraph()
    _add_horizontal_rule(hr)

    doc.add_paragraph()

    # ── Tender title ──────────────────────────────────────────────
    title_para = doc.add_heading(tender.get("title", "Untitled Tender"), level=1)
    title_para.runs[0].font.color.rgb = BRAND_DARK_BLUE

    doc.add_paragraph()

    # ── Key facts table ───────────────────────────────────────────
    doc.add_heading("📋 Tender Overview", level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    _style_table(table)

    rows_data = [
        ("Issuing Authority", tender.get("issuing_authority", "N/A")),
        ("Reference Number", tender.get("reference_number", "N/A")),
        ("Sector", tender.get("sector_name", "N/A")),
        ("Estimated Value", _format_value(tender.get("estimated_value_pkr"))),
        ("Published Date", _format_date(tender.get("published_date"))),
        ("Submission Deadline", _format_date(tender.get("submission_deadline"))),
        ("Opportunity Score", f"{tender.get('opportunity_score', 'N/A')}/100"),
    ]
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # ── AI Summary ────────────────────────────────────────────────
    doc.add_heading("🤖 AI-Generated Executive Summary", level=2)
    doc.add_paragraph(tender.get("ai_summary_short", "Summary not available."))
    doc.add_paragraph()

    # ── Detailed Analysis ─────────────────────────────────────────
    doc.add_heading("📊 Detailed Analysis", level=2)
    detailed = tender.get("ai_summary_detailed", "Detailed analysis not available.")
    for para in detailed.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph()

    # ── Eligibility ───────────────────────────────────────────────
    doc.add_heading("✅ Eligibility Criteria", level=2)
    eligibility = tender.get("eligibility_criteria", "Not specified.")
    doc.add_paragraph(eligibility)
    doc.add_paragraph()

    # ── SME Insights ─────────────────────────────────────────────
    doc.add_heading("💡 SME Insights & Recommendations", level=2)
    insights = tender.get("sme_insights", "")
    if insights:
        for line in insights.split("\n"):
            if line.strip():
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(line.strip().lstrip("•-– "))
    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────
    doc.add_paragraph()
    hr2 = doc.add_paragraph()
    _add_horizontal_rule(hr2)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        f"Generated by TenderIQ Pakistan on {datetime.utcnow().strftime('%d %B %Y at %H:%M UTC')} | "
        "www.tenderiq.pk | Data source: data.gov.pk (CKAN)"
    )
    fr.font.size = Pt(8)
    fr.font.italic = True

    # Serialize to bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() else "_" for c in tender.get("title", "tender")[:40])
    filename = f"TenderIQ_{safe_title}_{uuid.uuid4().hex[:6]}.docx"
    return buf.read(), filename


def _add_horizontal_rule(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0D47A1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _style_table(table):
    for row in table.rows:
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(4)


# ── PDF Document ──────────────────────────────────────────────────────────────

def generate_pdf_document(tender: Dict[str, Any]) -> Tuple[bytes, str]:
    """
    Generate a branded PDF document for a tender.
    Returns (bytes, filename).
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        leftMargin=1.0 * inch,
        rightMargin=1.0 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "TenderTitle",
        parent=styles["Title"],
        fontSize=18,
        textColor=RL_DARK_BLUE,
        spaceAfter=12,
        leading=22,
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=RL_DARK_BLUE,
        spaceBefore=16,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=15,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=2,
    )

    story = []

    # Header banner
    header_table = Table(
        [[
            Paragraph("<b><font color='#0D47A1' size='20'>🏛 TenderIQ Pakistan</font></b>", styles["Normal"]),
            Paragraph(
                f"<font color='#F57F17' size='9'><i>Government Procurement Intelligence</i></font>",
                styles["Normal"],
            ),
        ]],
        colWidths=[4 * inch, 2.5 * inch],
    )
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), RL_LIGHT_GREY),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [RL_LIGHT_GREY]),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.15 * inch))
    story.append(HRFlowable(width="100%", thickness=2, color=RL_DARK_BLUE))
    story.append(Spacer(1, 0.2 * inch))

    # Title
    story.append(Paragraph(tender.get("title", "Untitled Tender"), title_style))

    # Overview table
    story.append(Paragraph("📋 Tender Overview", h2_style))
    overview_data = [
        ["Field", "Details"],
        ["Issuing Authority", tender.get("issuing_authority", "N/A")],
        ["Reference Number", tender.get("reference_number", "N/A")],
        ["Sector", tender.get("sector_name", "N/A")],
        ["Estimated Value", _format_value(tender.get("estimated_value_pkr"))],
        ["Published Date", _format_date(tender.get("published_date"))],
        ["Submission Deadline", _format_date(tender.get("submission_deadline"))],
        ["Opportunity Score", f"{tender.get('opportunity_score', 'N/A')}/100"],
    ]
    ov_table = Table(overview_data, colWidths=[2 * inch, 4.5 * inch])
    ov_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), RL_DARK_BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, RL_LIGHT_GREY]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
    ]))
    story.append(ov_table)
    story.append(Spacer(1, 0.2 * inch))

    # AI Summary
    story.append(Paragraph("🤖 AI-Generated Executive Summary", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_GOLD))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        tender.get("ai_summary_short", "Summary not available."), body_style
    ))

    # Detailed analysis
    story.append(Paragraph("📊 Detailed Analysis", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_GOLD))
    story.append(Spacer(1, 0.1 * inch))
    detailed = tender.get("ai_summary_detailed", "")
    for para in detailed.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), body_style))

    # Eligibility
    story.append(Paragraph("✅ Eligibility Criteria", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_GOLD))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(tender.get("eligibility_criteria", "Not specified."), body_style))

    # SME Insights
    story.append(Paragraph("💡 SME Insights & Recommendations", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_GOLD))
    story.append(Spacer(1, 0.1 * inch))
    insights = tender.get("sme_insights", "")
    if insights:
        for line in insights.split("\n"):
            if line.strip():
                bullet = f"• {line.strip().lstrip('•-– ')}"
                story.append(Paragraph(bullet, body_style))

    # Footer
    story.append(Spacer(1, 0.3 * inch))
    story.append(HRFlowable(width="100%", thickness=1, color=RL_DARK_BLUE))
    story.append(Spacer(1, 0.1 * inch))
    footer_style = ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontSize=7, textColor=colors.grey, alignment=TA_CENTER,
    )
    story.append(Paragraph(
        f"Generated by TenderIQ Pakistan on "
        f"{datetime.utcnow().strftime('%d %B %Y at %H:%M UTC')} | "
        "www.tenderiq.pk | Data source: data.gov.pk (CKAN) | "
        "For informational purposes only.",
        footer_style,
    ))

    doc.build(story)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() else "_" for c in tender.get("title", "tender")[:40])
    filename = f"TenderIQ_{safe_title}_{uuid.uuid4().hex[:6]}.pdf"
    return buf.read(), filename
